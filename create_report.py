import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_code_block(doc, code_text):
    # Membuat tabel 1x1 sebagai wadah kotak kode
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run(code_text)
    # Font khusus kode (Monospace)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Jarak
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

def create_report():
    doc = Document()
    
    # Setting Kertas A4
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    
    # Margin 3 cm
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    
    # Setting Font Default
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    
    # 0. Identitas
    doc.add_paragraph('Nama: Intan Bella Safira Putri Dewi')
    doc.add_paragraph('Kelas: III Rekayasa Perangkat Lunak Kripto')
    doc.add_paragraph('NPM: 2322101923')
    doc.add_paragraph()
    
    # Judul
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run('LAPORAN TUGAS PRAKTIKUM\nPENGUJIAN PERANGKAT LUNAK (LOGIN & REGISTER)')
    run_title.bold = True
    doc.add_paragraph()
    
    # 1. Pendahuluan
    p_h1 = doc.add_paragraph()
    run_h1 = p_h1.add_run('1. Pendahuluan')
    run_h1.bold = True
    doc.add_paragraph('Laporan ini membahas pengujian fungsionalitas modul Login dan Register pada aplikasi "quiz-pengupil". Pengujian dilakukan dengan pendekatan Test Case skenario Positif dan Negatif, otomasi menggunakan framework Selenium dengan bahasa Python, dan implementasi pipeline CI/CD menggunakan GitHub Actions dengan menerapkan teknik Stub pada konfigurasi database.')
    doc.add_paragraph()

    # 2. Test Case (Berupa Tabel)
    p_h2 = doc.add_paragraph()
    run_h2 = p_h2.add_run('2. Skenario Pengujian (Test Case)')
    run_h2.bold = True
    doc.add_paragraph('Berdasarkan analisis kebutuhan sistem, dirancang 8 Test Case (4 untuk Register, 4 untuk Login) yang mendokumentasikan skenario input valid dan tidak valid. Berikut adalah tabel skenario tersebut:')
    
    # Membaca testcases.md dan mengubahnya menjadi Tabel asli di MS Word
    with open('tests/testcases.md', 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_data = []
    
    for line in lines:
        line = line.strip()
        if line.startswith('##'):
            if in_table and table_data:
                table = doc.add_table(rows=1, cols=len(table_data[0]))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, text in enumerate(table_data[0]):
                    hdr_cells[i].text = text
                    hdr_cells[i].paragraphs[0].runs[0].bold = True
                for row_data in table_data[2:]: # skip separator line
                    row_cells = table.add_row().cells
                    for i, text in enumerate(row_data):
                        text = text.replace('<br>', '\n')
                        row_cells[i].text = text
                doc.add_paragraph()
                in_table = False
                table_data = []
            
            p_sub = doc.add_paragraph()
            run_sub = p_sub.add_run(line.replace('##', '').strip())
            run_sub.bold = True
        elif line.startswith('|'):
            in_table = True
            row = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(row)
            
    if in_table and table_data:
        table = doc.add_table(rows=1, cols=len(table_data[0]))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, text in enumerate(table_data[0]):
            hdr_cells[i].text = text
            if hdr_cells[i].paragraphs[0].runs:
                hdr_cells[i].paragraphs[0].runs[0].bold = True
        for row_data in table_data[2:]:
            row_cells = table.add_row().cells
            for i, text in enumerate(row_data):
                text = text.replace('<br>', '\n')
                row_cells[i].text = text
        doc.add_paragraph()

    # 3. Implementasi Selenium
    p_h3 = doc.add_paragraph()
    run_h3 = p_h3.add_run('3. Implementasi Otomasi dengan Selenium')
    run_h3.bold = True
    doc.add_paragraph('Otomasi pengujian dilakukan menggunakan Python dan Selenium WebDriver. Script dikonfigurasi untuk dapat berjalan dengan GUI secara lokal maupun secara Headless (latar belakang tanpa tampilan) ketika ditarik oleh server CI/CD. Berikut adalah source code dari otomasi pengujian beserta hasil pengujian lokal:')
    
    with open('tests/test_auth.py', 'r', encoding='utf-8') as f:
        code_py = f.read()
    add_code_block(doc, code_py)
    
    doc.add_paragraph('Hasil eksekusi script otomasi pada server lokal menunjukkan seluruh skenario berhasil dilewati dengan sempurna (Passed).')
    doc.add_paragraph('[MASUKKAN SCREENSHOT TERMINAL VSCODE "Ran 8 tests in ... OK" DI SINI]')
    doc.add_paragraph()

    # 4. CI/CD dan Stub
    p_h4 = doc.add_paragraph()
    run_h4 = p_h4.add_run('4. Implementasi CI/CD dan Stubbing')
    run_h4.bold = True
    doc.add_paragraph('Automated test diintegrasikan dengan GitHub Actions. Mengingat GitHub menggunakan environment server bayangan (container), penerapan teknik Stub dilakukan pada script pipeline untuk menimpa kredensial koneksi database (koneksi.php) secara otomatis tanpa mengubah source code asli. Hal ini dicapai menggunakan perintah manipulasi teks "sed".')
    
    with open('.github/workflows/ci.yml', 'r', encoding='utf-8') as f:
        code_yml = f.read()
    add_code_block(doc, code_yml)
    doc.add_paragraph()

    # 5. Laporan Kendala Github
    p_h5 = doc.add_paragraph()
    run_h5 = p_h5.add_run('5. Laporan Kendala Pipeline CI/CD (Eksternal)')
    run_h5.bold = True
    doc.add_paragraph('Seluruh konfigurasi CI/CD (ci.yml) telah di-push dan sukses dieksekusi di localhost. Namun, eksekusi GitHub Actions terhambat di status "Queued". Hal ini murni dikarenakan server global layanan GitHub sedang mengalami status "Degraded Performance" (gangguan fungsionalitas Action dan antrean penuh). Berikut adalah lampiran bukti status dari halaman resmi GitHub Status yang diambil pada saat pengerjaan praktikum:')
    
    doc.add_paragraph('[MASUKKAN SCREENSHOT GITHUB STATUS DEGRADED DI SINI]')
    doc.add_paragraph('[MASUKKAN SCREENSHOT GITHUB ACTIONS QUEUED DI SINI]')
    
    doc.add_paragraph()
    doc.add_paragraph('Kesimpulan: Kode otomasi, skenario, dan pipeline yang dibuat telah 100% valid secara logika maupun teknis. Keberhasilan pengujian di local environment menjadi bukti keandalan test case dari kedua modul.')

    doc.save('2322101923_Intan Bella Safira Putri Dewi_Kuis.docx')
    print("Dokumen berhasil diperbarui dengan tabel dan kode box khusus!")

if __name__ == '__main__':
    create_report()
